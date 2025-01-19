import os
import openai
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# מפתח ה-API של OpenAI
openai.api_key = "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"

# פונקציה ליצירת תוכן מבוסס AI
def generate_content(prompt):
    response = openai.Completion.create(
        model="text-davinci-003",
        prompt=prompt,
        max_tokens=500
    )
    return response.choices[0].text.strip()

# פונקציה ליצירת קובץ Word
def create_word(content):
    doc = Document()
    doc.add_paragraph(content)
    doc.save("output.docx")

# פונקציה ליצירת קובץ PDF
def create_pdf(content):
    c = canvas.Canvas("output.pdf", pagesize=letter)
    c.drawString(100, 750, content)
    c.save()

# פונקציה ליצירת קובץ טקסט
def create_text(content):
    with open("output.txt", "w", encoding="utf-8") as file:
        file.write(content)

# פונקציה ליצירת קוד בשפת תכנות
def generate_code(prompt, language="Python"):
    response = openai.Completion.create(
        model="text-davinci-003",
        prompt=f"Write a {language} code that does the following: {prompt}",
        max_tokens=500
    )
    return response.choices[0].text.strip()

# תפריט ראשי
def main():
    print("ברוך הבא לכלי יצירת תוכן וקוד מבוסס AI!")
    print("בחר מה ברצונך לעשות:")
    print("1. ליצור תוכן")
    print("2. ליצור קוד")
    print("3. יציאה")

    choice = input("הזן את בחירתך (1/2/3): ")

    if choice == "1":
        prompt = input("הזן את ההנחיות ליצירת תוכן: ")
        content = generate_content(prompt)
        print("\nתוכן שנוצר:")
        print(content)

        print("\nבחר פורמט יצוא:")
        print("1. Word")
        print("2. PDF")
        print("3. טקסט רגיל")

        format_choice = input("הזן את בחירתך (1/2/3): ")
        if format_choice == "1":
            create_word(content)
            print("התוכן נשמר בקובץ output.docx")
        elif format_choice == "2":
            create_pdf(content)
            print("התוכן נשמר בקובץ output.pdf")
        elif format_choice == "3":
            create_text(content)
            print("התוכן נשמר בקובץ output.txt")
        else:
            print("בחירה לא חוקית.")
    elif choice == "2":
        prompt = input("הזן את ההנחיות ליצירת קוד: ")
        language = input("בחר שפת תכנות (Python, JavaScript, C++, וכו'): ")
        code = generate_code(prompt, language)
        print("\nקוד שנוצר:")
        print(code)

        save_choice = input("\nהאם לשמור את הקוד לקובץ? (y/n): ")
        if save_choice.lower() == "y":
            with open(f"code_output.{language.lower()}.txt", "w", encoding="utf-8") as file:
                file.write(code)
            print(f"הקוד נשמר בקובץ code_output.{language.lower()}.txt")
    elif choice == "3":
        print("להתראות!")
    else:
        print("בחירה לא חוקית.")

if __name__ == "__main__":
    main()
import os
import openai
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# מפתח ה-API של OpenAI
openai.api_key = "sk-xxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxx"

# פונקציה ליצירת תוכן מבוסס AI
def generate_content(prompt):
    response = openai.Completion.create(
        model="text-davinci-003",
        prompt=prompt,
        max_tokens=500
    )
    return response.choices[0].text.strip()

# פונקציה ליצירת קובץ Word
def create_word(content):
    doc = Document()
    doc.add_paragraph(content)
    doc.save("output.docx")

# פונקציה ליצירת קובץ PDF
def create_pdf(content):
    c = canvas.Canvas("output.pdf", pagesize=letter)
    c.drawString(100, 750, content)
    c.save()

# פונקציה ליצירת קובץ טקסט
def create_text(content):
    with open("output.txt", "w", encoding="utf-8") as file:
        file.write(content)

# פונקציה ליצירת קוד בשפת תכנות
def generate_code(prompt, language="Python"):
    response = openai.Completion.create(
        model="text-davinci-003",
        prompt=f"Write a {language} code that does the following: {prompt}",
        max_tokens=500
    )
    return response.choices[0].text.strip()

# תפריט ראשי
def main():
    print("ברוך הבא לכלי יצירת תוכן וקוד מבוסס AI!")
    print("בחר מה ברצונך לעשות:")
    print("1. ליצור תוכן")
    print("2. ליצור קוד")
    print("3. יציאה")

    choice = input("הזן את בחירתך (1/2/3): ")

    if choice == "1":
        prompt = input("הזן את ההנחיות ליצירת תוכן: ")
        content = generate_content(prompt)
        print("\nתוכן שנוצר:")
        print(content)

        print("\nבחר פורמט יצוא:")
        print("1. Word")
        print("2. PDF")
        print("3. טקסט רגיל")

        format_choice = input("הזן את בחירתך (1/2/3): ")
        if format_choice == "1":
            create_word(content)
            print("התוכן נשמר בקובץ output.docx")
        elif format_choice == "2":
            create_pdf(content)
            print("התוכן נשמר בקובץ output.pdf")
        elif format_choice == "3":
            create_text(content)
            print("התוכן נשמר בקובץ output.txt")
        else:
            print("בחירה לא חוקית.")
    elif choice == "2":
        prompt = input("הזן את ההנחיות ליצירת קוד: ")
        language = input("בחר שפת תכנות (Python, JavaScript, C++, וכו'): ")
        code = generate_code(prompt, language)
        print("\nקוד שנוצר:")
        print(code)

        save_choice = input("\nהאם לשמור את הקוד לקובץ? (y/n): ")
        if save_choice.lower() == "y":
            with open(f"code_output.{language.lower()}.txt", "w", encoding="utf-8") as file:
                file.write(code)
            print(f"הקוד נשמר בקובץ code_output.{language.lower()}.txt")
    elif choice == "3":
        print("להתראות!")
    else:
        print("בחירה לא חוקית.")

if __name__ == "__main__":
    main()
import os
import openai
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# מפתח ה-API של OpenAI
openai.api_key = "sk-xxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxx"

# פונקציה ליצירת תוכן מבוסס AI
def generate_content(prompt):
    response = openai.Completion.create(
        model="text-davinci-003",
        prompt=prompt,
        max_tokens=500
    )
    return response.choices[0].text.strip()

# פונקציה ליצירת קובץ Word
def create_word(content):
    doc = Document()
    doc.add_paragraph(content)
    doc.save("output.docx")

# פונקציה ליצירת קובץ PDF
def create_pdf(content):
    c = canvas.Canvas("output.pdf", pagesize=letter)
    c.drawString(100, 750, content)
    c.save()

# פונקציה ליצירת קובץ טקסט
def create_text(content):
    with open("output.txt", "w", encoding="utf-8") as file:
        file.write(content)

# פונקציה ליצירת קוד בשפת תכנות
def generate_code(prompt, language="Python"):
    response = openai.Completion.create(
        model="text-davinci-003",
        prompt=f"Write a {language} code that does the following: {prompt}",
        max_tokens=500
    )
    return response.choices[0].text.strip()

# תפריט ראשי
def main():
    print("ברוך הבא לכלי יצירת תוכן וקוד מבוסס AI!")
    print("בחר מה ברצונך לעשות:")
    print("1. ליצור תוכן")
    print("2. ליצור קוד")
    print("3. יציאה")

    choice = input("הזן את בחירתך (1/2/3): ")

    if choice == "1":
        prompt = input("הזן את ההנחיות ליצירת תוכן: ")
        content = generate_content(prompt)
        print("\nתוכן שנוצר:")
        print(content)

        print("\nבחר פורמט יצוא:")
        print("1. Word")
        print("2. PDF")
        print("3. טקסט רגיל")

        format_choice = input("הזן את בחירתך (1/2/3): ")
        if format_choice == "1":
            create_word(content)
            print("התוכן נשמר בקובץ output.docx")
        elif format_choice == "2":
            create_pdf(content)
            print("התוכן נשמר בקובץ output.pdf")
        elif format_choice == "3":
            create_text(content)
            print("התוכן נשמר בקובץ output.txt")
        else:
            print("בחירה לא חוקית.")
    elif choice == "2":
        prompt = input("הזן את ההנחיות ליצירת קוד: ")
        language = input("בחר שפת תכנות (Python, JavaScript, C++, וכו'): ")
        code = generate_code(prompt, language)
        print("\nקוד שנוצר:")
        print(code)

        save_choice = input("\nהאם לשמור את הקוד לקובץ? (y/n): ")
        if save_choice.lower() == "y":
            with open(f"code_output.{language.lower()}.txt", "w", encoding="utf-8") as file:
                file.write(code)
            print(f"הקוד נשמר בקובץ code_output.{language.lower()}.txt")
    elif choice == "3":
        print("להתראות!")
    else:
        print("בחירה לא חוקית.")

if __name__ == "__main__":
    main()
